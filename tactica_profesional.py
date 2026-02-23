# tactica_profesional.py
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# TÃCTICA INGENIERÃA â€” Backend Profesional v3.1 (fix static + rutas)
# FastAPI + GPT-4o Vision + Reportes PDF/Excel profesionales
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

import os
import io
import re
import json
import base64
import logging
import datetime
import sqlite3
import uuid
import traceback
from pathlib import Path
from typing import List, Optional, Dict

from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware

import uvicorn
import httpx
from PIL import Image as PILImage

# PyMuPDF para PDFs
try:
    import fitz
except ImportError:
    fitz = None

# OpenAI
try:
    import openai
    from openai import (
        APIConnectionError as OAIConnectionError,
        APITimeoutError as OAITimeoutError,
        RateLimitError as OAIRateLimitError,
        APIStatusError as OAIStatusError,
        AuthenticationError as OAIAuthError,
    )
except ImportError:
    openai = None
    OAIConnectionError = OAITimeoutError = OAIRateLimitError = OAIStatusError = OAIAuthError = Exception

# Pandas
try:
    import pandas as pd
except ImportError:
    pd = None

# ReportLab (PDF profesional)
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.colors import grey, white
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.colors import HexColor
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image as RLImage
    )
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

# Openpyxl (Excel profesional)
try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.chart import BarChart, Reference
    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False


# python-docx (Word profesional)
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# Matplotlib para grÃ¡ficos estilo BI
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    MATPLOTLIB_OK = True
except ImportError:
    MATPLOTLIB_OK = False


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  CONFIGURACIÃ“N
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

BASE_DIR = Path(__file__).resolve().parent
STATIC_DIR = BASE_DIR / "static"
ASSETS_DIR = STATIC_DIR / "assets"
EXPORTS_DIR = BASE_DIR / "exports"
OUTPUT_DIR = EXPORTS_DIR
UPLOAD_DIR = STATIC_DIR / "uploads"

for d in (STATIC_DIR, ASSETS_DIR, EXPORTS_DIR, OUTPUT_DIR, UPLOAD_DIR):
    d.mkdir(parents=True, exist_ok=True)

LOGO_FILENAME = "LOGO TACTICA INGENIERIA.png"
LOGO_PATH = ASSETS_DIR / LOGO_FILENAME
REPORT_VERSION = "3.2.0"

def _resolve_logo_path() -> Optional[Path]:
    """Busca un logo en static/assets (prioriza LOGO_PATH, luego pngs)."""
    if LOGO_PATH.exists():
        return LOGO_PATH
    if ASSETS_DIR.exists():
        for p in ASSETS_DIR.glob("*.png"):
            return p
    return None

def _report_timestamp() -> str:
    # Formato requerido: YYYYMMDD_HHMM
    return datetime.datetime.now().strftime("%Y%m%d_%H%M")

def _report_filename(ts: str, ext: str) -> str:
    base = f"Informe_{ts}"
    ext = ext.lstrip(".")
    name = f"{base}.{ext}"
    # Evitar colisiones en el mismo minuto
    if (EXPORTS_DIR / name).exists():
        for i in range(2, 10):
            name = f"{base}_{i}.{ext}"
            if not (EXPORTS_DIR / name).exists():
                break
    return name

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  BASE DE DATOS LOCAL (SQLite) â€” NÃºcleo del Proyecto
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

DB_PATH = BASE_DIR / "tactica.db"

def _db_connect():
    conn = sqlite3.connect(str(DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn

def _db_init():
    conn = _db_connect()
    cur = conn.cursor()
    # âœ… FIX: desactivar FK enforcement para que project_id NULL funcione sin error
    cur.execute("PRAGMA foreign_keys = OFF")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS projects (
        id TEXT PRIMARY KEY,
        name TEXT NOT NULL,
        location TEXT,
        currency TEXT DEFAULT 'COP',
        budget_total REAL,
        created_at TEXT NOT NULL,
        updated_at TEXT NOT NULL
    )
    """)

    # âœ… FIX: tabla events SIN FOREIGN KEY constraint para permitir project_id NULL
    cur.execute("""
    CREATE TABLE IF NOT EXISTS events (
        id TEXT PRIMARY KEY,
        project_id TEXT,
        type TEXT NOT NULL,
        payload_json TEXT NOT NULL,
        created_at TEXT NOT NULL
    )
    """)
    # MigraciÃ³n automÃ¡tica si la tabla vieja tiene FK constraint
    try:
        cur.execute("INSERT INTO events(id,project_id,type,payload_json,created_at) VALUES('__fk_test__',NULL,'__test__','{}','2000-01-01T00:00:00Z')")
        cur.execute("DELETE FROM events WHERE id='__fk_test__'")
    except Exception:
        # La tabla existe con FK â†’ migrar sin FK
        cur.execute("ALTER TABLE events RENAME TO events_old")
        cur.execute("""
        CREATE TABLE events (
            id TEXT PRIMARY KEY,
            project_id TEXT,
            type TEXT NOT NULL,
            payload_json TEXT NOT NULL,
            created_at TEXT NOT NULL
        )
        """)
        cur.execute("INSERT OR IGNORE INTO events SELECT id,project_id,type,payload_json,created_at FROM events_old")
        cur.execute("DROP TABLE events_old")

    cur.execute("""
    CREATE TABLE IF NOT EXISTS reports (
        id TEXT PRIMARY KEY,
        project_id TEXT,
        module TEXT NOT NULL,
        title TEXT NOT NULL,
        subtitle TEXT,
        files_json TEXT NOT NULL,
        source_event_id TEXT,
        created_at TEXT NOT NULL,
        FOREIGN KEY(project_id) REFERENCES projects(id)
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS settings (
        key TEXT PRIMARY KEY,
        value TEXT NOT NULL
    )
    """)

    # defaults branding
    cur.execute("INSERT OR IGNORE INTO settings(key,value) VALUES('company_name', ?)", ("TÃCTICA INGENIERÃA",))
    cur.execute("INSERT OR IGNORE INTO settings(key,value) VALUES('logo_filename', ?)", (LOGO_FILENAME,))
    conn.commit()
    conn.close()

def _now_iso():
    return datetime.datetime.utcnow().replace(microsecond=0).isoformat() + "Z"

def _new_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:16]}"

def _db_insert_event(project_id: Optional[str], type_: str, payload: dict) -> str:
    """Inserta un evento. project_id puede ser None (sin proyecto activo)."""
    eid = _new_id("evt")
    # Normalizar: string vacÃ­o â†’ None
    pid = project_id if project_id and str(project_id).strip() else None
    conn = _db_connect()
    # âœ… FIX: desactivar FK enforcement por si la tabla vieja tiene constraint
    conn.execute("PRAGMA foreign_keys = OFF")
    conn.execute(
        "INSERT INTO events(id, project_id, type, payload_json, created_at) VALUES(?,?,?,?,?)",
        (eid, pid, type_, json.dumps(payload, ensure_ascii=False, default=str), _now_iso())
    )
    conn.commit()
    conn.close()
    return eid

def _db_get_event(event_id: str) -> Optional[dict]:
    conn = _db_connect()
    row = conn.execute("SELECT * FROM events WHERE id=?", (event_id,)).fetchone()
    conn.close()
    if not row:
        return None
    return {
        "id": row["id"],
        "project_id": row["project_id"],
        "type": row["type"],
        "payload": json.loads(row["payload_json"]),
        "created_at": row["created_at"],
    }

def _db_get_setting(key: str, default: str = "") -> str:
    conn = _db_connect()
    row = conn.execute("SELECT value FROM settings WHERE key=?", (key,)).fetchone()
    conn.close()
    return (row["value"] if row else default)

def _get_project_meta(project_id: Optional[str]) -> dict:
    if not project_id:
        return {}
    conn = _db_connect()
    row = conn.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
    conn.close()
    if not row:
        return {}
    return {
        "project_id": row["id"],
        "project_name": row["name"],
        "project_location": row["location"] or "",
        "currency": row["currency"] or "",
        "company_name": _db_get_setting("company_name", "TÁCTICA INGENIERÍA"),
    }


OPENAI_API_KEY = (os.getenv("OPENAI_API_KEY") or "").strip()
OPENAI_MODEL   = os.getenv("OPENAI_MODEL", "gpt-4o")   # configurable sin tocar código
SIMULATION = (not OPENAI_API_KEY) or (OPENAI_API_KEY.upper() == "SIMULACION")

# Colores corporativos
CORP_TEAL = "#4db8a4"
CORP_DARK = "#060b10"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(name)s] %(levelname)s: %(message)s"
)
logger = logging.getLogger("tactica")


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  FASTAPI APP
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

app = FastAPI(
    title="TÃ¡ctica IngenierÃ­a API",
    description="Backend profesional para mediciÃ³n, anÃ¡lisis estructural y presupuestos",
    version="3.1.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# âœ… FIX CLAVE: html=True ayuda con /static/ (index) y rutas de html amigables
# âœ… /static: lo servimos nosotros (con fallback), para matar el 404 incluso si algÃºn archivo quedÃ³ como *.html.html
app.mount("/outputs", StaticFiles(directory=str(EXPORTS_DIR), html=True), name="outputs")
app.mount("/exports", StaticFiles(directory=str(EXPORTS_DIR), html=True), name="exports")
app.mount("/uploads", StaticFiles(directory=str(UPLOAD_DIR), html=True), name="uploads")

# Ãndice rÃ¡pido de archivos estÃ¡ticos (para resolver nombres raros tipo index.html.html)
STATIC_INDEX: Dict[str, Path] = {}

# Alias por si el front o tÃº escriben nombres â€œhumanosâ€
STATIC_ALIASES = {
    "estructural.html": "structural.html",
    "budget.html": "budget_advanced.html",
    "dashboard.html": "dashboard_panel.html",
}

def _build_static_index() -> None:
    STATIC_INDEX.clear()
    if not STATIC_DIR.exists():
        return
    for root, _, files in os.walk(STATIC_DIR):
        for f in files:
            rel = Path(root).relative_to(STATIC_DIR) / f
            key = str(rel).replace("\\", "/").lower()
            STATIC_INDEX[key] = rel

def _resolve_static(rel_path: str) -> Optional[Path]:
    # Normaliza
    rel_path = (rel_path or "").lstrip("/")
    if rel_path == "" or rel_path.endswith("/"):
        rel_path = rel_path + "index.html"

    # Alias directos
    if rel_path in STATIC_ALIASES:
        rel_path = STATIC_ALIASES[rel_path]

    # Candidatos (en orden)
    candidates = [rel_path]

    lp = rel_path.lower()

    # Si pidieron .html, prueba tambiÃ©n .html.html (error tÃ­pico de Windows al â€œguardar comoâ€)
    if lp.endswith(".html"):
        candidates.append(rel_path + ".html")
    # Si pidieron sin extensiÃ³n, prueba .html y .html.html
    if "." not in Path(rel_path).name:
        candidates.append(rel_path + ".html")
        candidates.append(rel_path + ".html.html")
    # Si pidieron .html.html, intenta la versiÃ³n normal
    if lp.endswith(".html.html"):
        candidates.append(rel_path[:-5])  # quita el Ãºltimo ".html"

    # Aplicar alias tambiÃ©n a candidatos
    extra = []
    for c in list(candidates):
        if c in STATIC_ALIASES:
            extra.append(STATIC_ALIASES[c])
    candidates.extend(extra)

    # Buscar en el Ã­ndice (case-insensitive)
    for c in candidates:
        key = c.replace("\\", "/").lower()
        hit = STATIC_INDEX.get(key)
        if hit is not None:
            return hit

    # Si no se encontrÃ³, re-indexa 1 vez (por si acabas de copiar archivos)
    _build_static_index()
    for c in candidates:
        key = c.replace("\\", "/").lower()
        hit = STATIC_INDEX.get(key)
        if hit is not None:
            return hit

    return None

@app.on_event("startup")
async def _startup_build_index():
    _build_static_index()
    _db_init()

@app.get("/static", include_in_schema=False)
async def static_root():
    return RedirectResponse("/static/index.html")

@app.get("/static/{path:path}", include_in_schema=False)
async def static_files(path: str):
    rel = _resolve_static(path)
    if rel is None:
        raise HTTPException(404, f"Archivo estÃ¡tico no encontrado: {path}")
    full = (STATIC_DIR / rel).resolve()
    # Seguridad anti traversal
    if not str(full).startswith(str(STATIC_DIR.resolve())):
        raise HTTPException(400, "Ruta invÃ¡lida")
    return FileResponse(str(full))


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  PROMPTS DE IA (GPT-4o Vision)
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

SYSTEM_PROMPT_BASE = """Eres un ingeniero civil senior experto en anÃ¡lisis de planos estructurales,
mediciones de obra, cartillas de hierro, presupuestos de construcciÃ³n y normativa NSR-10 (Colombia).
Respondes siempre en espaÃ±ol. Eres preciso, tÃ©cnico y profesional.
Cuando analices imÃ¡genes de planos, sÃ© extremadamente detallado con medidas, escalas y elementos."""

PROMPT_MEDICION = """Analiza este plano/imagen de ingenierÃ­a civil. Extrae TODAS las medidas visibles.
Para cada elemento encontrado, indica:
- elemento: nombre del elemento (viga, columna, losa, aleta, muro, etc.)
- medida: valor numÃ©rico de la medida
- unidad: unidad (m, cm, mm, mÂ², mÂ³)
- observacion: cualquier detalle relevante (ubicaciÃ³n, tipo, especificaciÃ³n)

Responde ÃšNICAMENTE con un JSON array vÃ¡lido. Ejemplo:
[{"elemento":"Viga V-1","medida":6.50,"unidad":"m","observacion":"Eje A-B, piso 2"}]"""

PROMPT_ESTRUCTURAL = """Analiza este plano estructural en detalle. Identifica TODOS los elementos estructurales:
vigas, columnas, zapatas, muros, losas, escaleras, etc.

Para cada elemento indica:
- elemento: identificaciÃ³n completa (ej: "Columna C-1", "Viga V-201")
- tipo: tipo estructural (columna, viga, zapata, losa, muro, etc.)
- seccion: dimensiones de la secciÃ³n (ej: "30x40 cm", "âˆ…60 cm")
- longitud_m: longitud estimada en metros
- area_m2: Ã¡rea de la secciÃ³n en mÂ² (si aplica)
- refuerzo: descripciÃ³n del refuerzo visible (si se ve)
- ubicacion: posiciÃ³n en el plano (ejes, nivel, piso)
- observacion: notas adicionales

Responde ÃšNICAMENTE con un JSON array vÃ¡lido."""

PROMPT_CARTILLA_HIERRO = """Analiza esta cartilla/despiece de hierro (rebar schedule). Extrae CADA barra/varilla:

Para cada elemento de refuerzo indica:
- elemento: elemento al que pertenece (ej: "Viga V-1", "Columna C-3")
- marca: marca o nÃºmero de la barra (ej: "1", "2a", "E-1")
- diametro_mm: diÃ¡metro en mm (ej: 12, 16, 20, 25)  â€” tambiÃ©n conocido como #3=9.5mm, #4=12.7mm, #5=15.9mm, #6=19.1mm, #7=22.2mm, #8=25.4mm
- cantidad: nÃºmero de barras
- longitud_m: longitud de cada barra en metros
- peso_unit_kg: peso unitario en kg (usar tabla: #3=0.56kg/m, #4=0.994kg/m, #5=1.552kg/m, #6=2.235kg/m, #7=3.042kg/m, #8=3.973kg/m)
- peso_total_kg: peso total = cantidad Ã— longitud_m Ã— peso_unit_kg
- forma: descripciÃ³n de la forma (recta, gancho, estribo, U, L, etc.)
- separacion_cm: separaciÃ³n si es estribo/refuerzo transversal
- observacion: notas

Responde ÃšNICAMENTE con un JSON array vÃ¡lido. Calcula los pesos con precisiÃ³n."""

PROMPT_PRESUPUESTO = """Analiza este documento de presupuesto/APU de construcciÃ³n. Extrae los items:

Para cada partida indica:
- item: cÃ³digo del item
- descripcion: descripciÃ³n de la actividad
- unidad: unidad de medida (m, mÂ², mÂ³, kg, un, gl)
- cantidad: cantidad presupuestada
- valor_unitario: precio unitario
- valor_total: subtotal de la partida
- capitulo: capÃ­tulo al que pertenece

Responde ÃšNICAMENTE con un JSON array vÃ¡lido."""

PROMPT_CHAT_TECNICO = """Eres el asistente tÃ©cnico de TÃ¡ctica IngenierÃ­a.
Responde consultas sobre: normativa estructural (NSR-10, ACI 318), diseÃ±o de concreto reforzado,
cuantÃ­as de acero, resistencia de materiales, cÃ¡lculo de cargas, interpretaciÃ³n de planos,
cartillas de hierro, presupuestos de obra, APUs, y todo lo relacionado con ingenierÃ­a civil.
Si el usuario envÃ­a una imagen, analÃ­zala en contexto de ingenierÃ­a civil.
SÃ© preciso, tÃ©cnico y Ãºtil. Responde en espaÃ±ol."""


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  UTILIDADES
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def get_openai_client():
    """
    Retorna cliente OpenAI con httpx configurado:
    - timeout de 30s en todas las operaciones
    - trust_env=True respeta HTTP_PROXY / HTTPS_PROXY del sistema (útil en redes corporativas)
    Retorna None si no hay key o si el paquete no está instalado.
    """
    if SIMULATION or not openai:
        return None
    _http = httpx.Client(timeout=httpx.Timeout(30.0, connect=10.0), trust_env=True)
    return openai.OpenAI(api_key=OPENAI_API_KEY, http_client=_http)


def image_to_base64(img: PILImage.Image, max_size: int = 2048) -> str:
    """Convierte PIL Image a base64 JPEG, redimensionando si es necesario."""
    if img.mode != "RGB":
        img = img.convert("RGB")
    w, h = img.size
    if max(w, h) > max_size:
        ratio = max_size / max(w, h)
        img = img.resize((int(w * ratio), int(h * ratio)), PILImage.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def file_to_images(contents: bytes, filename: str) -> List[PILImage.Image]:
    """Convierte archivo (imagen o PDF) a lista de PIL Images."""
    images: List[PILImage.Image] = []
    lower = (filename or "").lower()

    if lower.endswith(".pdf"):
        if fitz is None:
            raise HTTPException(500, "PyMuPDF no instalado. Instala: pip install PyMuPDF")
        doc = fitz.open(stream=contents, filetype="pdf")
        for page_num in range(min(doc.page_count, 5)):
            pix = doc.load_page(page_num).get_pixmap(matrix=fitz.Matrix(2, 2))
            img = PILImage.open(io.BytesIO(pix.tobytes("png")))
            images.append(img)
        doc.close()
    elif lower.endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp")):
        img = PILImage.open(io.BytesIO(contents))
        images.append(img)
    elif lower.endswith((".xlsx", ".xls")):
        # âœ… FIX: Excel no es imagen â†’ lo convertimos a imagen PNG via tabla
        if not OPENPYXL_OK:
            raise HTTPException(500, "openpyxl no instalado para leer Excel")
        data = parse_excel_file(contents, filename)
        if not data:
            raise HTTPException(400, "Excel vacÃ­o o sin datos")
        # Render tabla como imagen usando Pillow
        from PIL import ImageDraw, ImageFont
        cols = list(data[0].keys()) if data else []
        rows_data = [[str(r.get(c, "")) for c in cols] for r in data[:50]]
        col_w = 160
        row_h = 24
        padding = 10
        img_w = max(800, len(cols) * col_w + padding * 2)
        img_h = (len(rows_data) + 2) * row_h + padding * 2
        img = PILImage.new("RGB", (img_w, img_h), color=(15, 23, 42))
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("arial.ttf", 12)
        except Exception:
            font = ImageFont.load_default()
        # Header
        for ci, col in enumerate(cols):
            draw.rectangle([padding + ci*col_w, padding, padding + (ci+1)*col_w - 2, padding + row_h - 2], fill=(34, 193, 195, 180))
            draw.text((padding + ci*col_w + 4, padding + 4), str(col)[:18], font=font, fill=(255, 255, 255))
        # Rows
        for ri, row in enumerate(rows_data):
            bg = (20, 30, 50) if ri % 2 == 0 else (25, 38, 60)
            for ci, val in enumerate(row):
                y = padding + (ri + 1) * row_h
                draw.rectangle([padding + ci*col_w, y, padding + (ci+1)*col_w - 2, y + row_h - 2], fill=bg)
                draw.text((padding + ci*col_w + 4, y + 4), str(val)[:18], font=font, fill=(220, 230, 255))
        images.append(img)
    elif lower.endswith(".csv"):
        if not pd:
            raise HTTPException(500, "pandas no instalado para CSV")
        df = pd.read_csv(io.BytesIO(contents))
        data = df.head(50).to_dict(orient="records")
        if not data:
            raise HTTPException(400, "CSV vacÃ­o")
        from PIL import ImageDraw, ImageFont
        cols = list(data[0].keys())
        col_w = 160
        row_h = 24
        padding = 10
        img_w = max(800, len(cols)*col_w + padding*2)
        img_h = (len(data)+2)*row_h + padding*2
        img = PILImage.new("RGB", (img_w, img_h), color=(15, 23, 42))
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("arial.ttf", 12)
        except Exception:
            font = ImageFont.load_default()
        for ci, col in enumerate(cols):
            draw.rectangle([padding+ci*col_w, padding, padding+(ci+1)*col_w-2, padding+row_h-2], fill=(34,193,195))
            draw.text((padding+ci*col_w+4, padding+4), str(col)[:18], font=font, fill=(255,255,255))
        for ri, row in enumerate(data):
            bg = (20,30,50) if ri%2==0 else (25,38,60)
            for ci, col in enumerate(cols):
                y = padding+(ri+1)*row_h
                draw.rectangle([padding+ci*col_w, y, padding+(ci+1)*col_w-2, y+row_h-2], fill=bg)
                draw.text((padding+ci*col_w+4, y+4), str(row.get(col,""))[:18], font=font, fill=(220,230,255))
        images.append(img)
    else:
        raise HTTPException(400, f"Formato no soportado: {filename}. Usa PDF, imagen, Excel o CSV.")

    return images


def parse_excel_file(contents: bytes, filename: str) -> List[Dict]:
    """Lee un archivo Excel y retorna lista de dicts."""
    if not OPENPYXL_OK:
        raise HTTPException(500, "openpyxl no instalado")

    wb = load_workbook(io.BytesIO(contents), data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))

    if not rows:
        wb.close()
        return []

    header_idx = 0
    for i, row in enumerate(rows):
        if any(cell is not None for cell in row):
            header_idx = i
            break

    headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(rows[header_idx])]
    data: List[Dict] = []
    for row in rows[header_idx + 1:]:
        if any(cell is not None for cell in row):
            record = {}
            for j, cell in enumerate(row):
                if j < len(headers):
                    record[headers[j]] = cell
            data.append(record)

    wb.close()
    return data


def _summarize_items(items: List[dict]) -> dict:
    summary = {
        "count": len(items),
        "headers": [],
        "numeric": {},
    }
    if not items:
        return summary
    headers = []
    for it in items[:50]:
        for k in it.keys():
            if k not in headers:
                headers.append(k)
        if len(headers) >= 20:
            break
    summary["headers"] = headers
    for h in headers:
        vals = []
        for it in items:
            v = it.get(h)
            if isinstance(v, (int, float)):
                vals.append(float(v))
        if vals:
            summary["numeric"][h] = {
                "min": min(vals),
                "max": max(vals),
                "avg": sum(vals) / len(vals),
                "sum": sum(vals),
            }
    return summary


def _project_line(project_meta: Optional[dict]) -> str:
    if not project_meta:
        return ""
    name = project_meta.get("project_name") or ""
    loc = project_meta.get("project_location") or ""
    client = project_meta.get("client_name") or ""
    parts = [p for p in [name, loc, client] if p]
    return " | ".join(parts)


def call_gpt4o_vision(prompt: str, images_b64: List[str], system: str = SYSTEM_PROMPT_BASE) -> Optional[str]:
    """
    Llama a GPT-4o Vision con imágenes (usa chat.completions — la Responses API
    aún no soporta image_url como input directo).
    Errores granulares: 503 (conexión), 429 (rate limit), 504 (timeout), 401 (auth), 502 (error API).
    """
    client = get_openai_client()
    if client is None:
        return None

    content = [{"type": "text", "text": prompt}]
    for img_b64 in images_b64:
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/jpeg;base64,{img_b64}", "detail": "high"}
        })

    try:
        res = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": content}
            ],
            max_tokens=4096,
            temperature=0.1,
        )
        return res.choices[0].message.content

    except OAIAuthError:
        logger.error("OpenAI: clave API inválida (AuthenticationError)")
        raise HTTPException(401, "OPENAI_API_KEY inválida. Verifica tu clave en el archivo .env o en las variables de entorno.")

    except OAIConnectionError as e:
        logger.error(f"OpenAI: error de conexión (APIConnectionError): {e}")
        raise HTTPException(503, "No se pudo conectar con OpenAI. Revisa tu conexión a Internet, firewall o proxy. "
                                 "Si usas red corporativa, configura HTTPS_PROXY en tus variables de entorno.")

    except OAITimeoutError as e:
        logger.error(f"OpenAI: timeout (APITimeoutError): {e}")
        raise HTTPException(504, "Tiempo agotado al llamar a OpenAI (30s). El servidor está tardando — intenta de nuevo.")

    except OAIRateLimitError as e:
        logger.warning(f"OpenAI: rate limit alcanzado: {e}")
        raise HTTPException(429, "Límite de solicitudes GPT-4o alcanzado. Espera 30 segundos e intenta de nuevo.")

    except OAIStatusError as e:
        logger.error(f"OpenAI: error de estado HTTP {e.status_code}: {e}")
        raise HTTPException(502, f"OpenAI respondió con error {e.status_code}. Intenta en unos momentos.")

    except Exception as e:
        logger.error(f"GPT-4o vision error inesperado: {traceback.format_exc()}")
        raise HTTPException(500, f"Error interno al llamar a GPT-4o Vision: {str(e)}")


def call_gpt4o_text(prompt: str, system: str = SYSTEM_PROMPT_BASE) -> Optional[str]:
    """
    Llama a GPT-4o solo con texto usando la Responses API (client.responses.create).
    Misma gestión de errores que call_gpt4o_vision para consistencia.
    """
    client = get_openai_client()
    if client is None:
        return None

    try:
        resp = client.responses.create(
            model=OPENAI_MODEL,
            instructions=system,
            input=prompt,
        )
        return resp.output_text

    except OAIAuthError:
        logger.error("OpenAI: clave API inválida")
        raise HTTPException(401, "OPENAI_API_KEY inválida. Verifica tu clave en el archivo .env o en las variables de entorno.")

    except OAIConnectionError as e:
        logger.error(f"OpenAI: error de conexión: {e}")
        raise HTTPException(503, "No se pudo conectar con OpenAI. Revisa tu conexión a Internet, firewall o proxy. "
                                 "Si usas red corporativa, configura HTTPS_PROXY en tus variables de entorno.")

    except OAITimeoutError as e:
        logger.error(f"OpenAI: timeout: {e}")
        raise HTTPException(504, "Tiempo agotado al llamar a OpenAI (30s). Intenta de nuevo en unos momentos.")

    except OAIRateLimitError as e:
        logger.warning(f"OpenAI: rate limit: {e}")
        raise HTTPException(429, "Límite de solicitudes alcanzado. Espera 30 segundos e intenta de nuevo.")

    except OAIStatusError as e:
        logger.error(f"OpenAI: error HTTP {e.status_code}: {e}")
        raise HTTPException(502, f"OpenAI respondió con error {e.status_code}. Intenta en unos momentos.")

    except Exception as e:
        logger.error(f"GPT-4o text error inesperado: {traceback.format_exc()}")
        raise HTTPException(500, f"Error interno al llamar a GPT-4o: {str(e)}")


def parse_json_from_response(text: Optional[str]) -> List[dict]:
    """Extrae JSON array de la respuesta."""
    if not text:
        return []
    cleaned = text.strip()
    cleaned = re.sub(r"```json\s*", "", cleaned)
    cleaned = re.sub(r"```\s*", "", cleaned).strip()

    start = cleaned.find("[")
    end = cleaned.rfind("]")
    if start != -1 and end != -1:
        try:
            return json.loads(cleaned[start:end + 1])
        except json.JSONDecodeError:
            pass

    try:
        obj = json.loads(cleaned)
        return obj if isinstance(obj, list) else [obj]
    except json.JSONDecodeError:
        logger.warning(f"No se pudo parsear JSON: {cleaned[:200]}...")
        return []



def normalize_mode(mode: str) -> str:
    m = (mode or "measure").strip().lower()
    aliases = {
        "structural": "struct",
        "estructural": "struct",
        "structure": "struct",
        "medicion": "measure",
        "mediciÃ³n": "measure",
        "measurement": "measure",
        "rebar_schedule": "rebar",
        "cartilla": "rebar",
        "hierro": "rebar",
        "presupuesto": "budget",
        "costos": "budget",
    }
    return aliases.get(m, m)

def get_timestamp() -> str:
    return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  GENERACIÃ“N DE REPORTES PROFESIONALES
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def generar_reporte_pdf_profesional(
    items: List[dict],
    out_path: Path,
    titulo: str = "Reporte de Mediciones",
    subtitulo: str = "",
    columnas: Optional[List[str]] = None,
    logo_path: Optional[Path] = None,
    project_meta: Optional[dict] = None
):
    if not REPORTLAB_OK:
        logger.warning("ReportLab no disponible, PDF no generado")
        return

    if not logo_path:
        logo_path = _resolve_logo_path()

    width, height = letter
    ts = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    company = (project_meta or {}).get("company_name") or _db_get_setting("company_name", "TÁCTICA INGENIERÍA")
    project_line = _project_line(project_meta)
    summary = _summarize_items(items)

    def footer(c, doc_):
        c.saveState()
        c.setFillColor(HexColor("#777777"))
        c.setFont("Helvetica", 8)
        c.drawString(30, 20, f"Generado: {ts}")
        c.drawRightString(width - 30, 20, f"Página {c.getPageNumber()} | v{REPORT_VERSION}")
        c.restoreState()

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=letter,
        topMargin=50,
        bottomMargin=35,
        leftMargin=30,
        rightMargin=30,
    )

    styles = getSampleStyleSheet()
    styles["Title"].fontSize = 22
    styles["Title"].leading = 26
    styles["Heading2"].textColor = HexColor(CORP_TEAL)

    elements = []

    # Portada
    if logo_path and logo_path.exists():
        try:
            img = RLImage(str(logo_path), width=80, height=80)
            img.hAlign = "LEFT"
            elements.append(img)
        except Exception:
            pass
    elements.append(Spacer(1, 18))
    elements.append(Paragraph(company, styles["Title"]))
    elements.append(Paragraph(titulo, styles["Heading2"]))
    if subtitulo:
        elements.append(Paragraph(subtitulo, styles["Normal"]))
    if project_line:
        elements.append(Paragraph(f"Proyecto/Cliente: {project_line}", styles["Normal"]))
    elements.append(Paragraph(f"Fecha: {ts}", styles["Normal"]))
    elements.append(Paragraph(f"Versión del reporte: v{REPORT_VERSION}", styles["Normal"]))
    elements.append(PageBreak())

    # Resumen ejecutivo
    elements.append(Paragraph("Resumen ejecutivo", styles["Heading2"]))
    summary_lines = [f"Se procesaron {summary['count']} registros con {len(summary['headers'])} campos."]
    if summary["numeric"]:
        for k, v in list(summary["numeric"].items())[:3]:
            summary_lines.append(
                f"{k}: min {v['min']:.2f} | max {v['max']:.2f} | promedio {v['avg']:.2f}"
            )
    elements.append(Paragraph("<br/>".join(summary_lines), styles["Normal"]))
    elements.append(Spacer(1, 8))

    # Detalle
    elements.append(Paragraph("Detalle", styles["Heading2"]))
    if items:
        headers = columnas if columnas else summary["headers"] or list(items[0].keys())
        table_data = [headers]
        for item in items:
            table_data.append([str(item.get(h, "")) for h in headers])
        col_count = len(headers)
        col_width = (width - 60) / max(col_count, 1)
        table = Table(table_data, colWidths=[col_width] * col_count, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HexColor(CORP_TEAL)),
            ("TEXTCOLOR", (0, 0), (-1, 0), white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [white, HexColor("#f5f8fa")]),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        elements.append(table)
    else:
        elements.append(Paragraph("No se encontraron datos para mostrar.", styles["Normal"]))

    elements.append(Spacer(1, 12))
    elements.append(Paragraph("Conclusiones y recomendaciones", styles["Heading2"]))
    concl = [
        "Validar los registros clave y documentar supuestos de cálculo.",
        "Revisar los valores extremos detectados en las columnas numéricas.",
        "Consolidar esta versión como línea base del proyecto.",
    ]
    elements.append(Paragraph("<br/>".join([f"• {c}" for c in concl]), styles["Normal"]))

    doc.build(elements, onFirstPage=footer, onLaterPages=footer)
    logger.info(f"PDF generado: {out_path}")

def generar_reporte_excel_profesional(
    items: List[dict],
    out_path: Path,
    titulo: str = "Reporte de Mediciones",
    hoja: str = "Detalle",
    logo_path: Optional[Path] = None,
    incluir_graficos: bool = False,
    project_meta: Optional[dict] = None,
):
    if not OPENPYXL_OK:
        logger.warning("openpyxl no disponible, Excel no generado")
        return

    if not logo_path:
        logo_path = _resolve_logo_path()

    wb = Workbook()
    ws_summary = wb.active
    ws_summary.title = "Resumen"
    ws_detail = wb.create_sheet("Detalle")

    teal_fill = PatternFill(start_color="4DB8A4", end_color="4DB8A4", fill_type="solid")
    alt_fill = PatternFill(start_color="F5F8FA", end_color="F5F8FA", fill_type="solid")
    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
    title_font = Font(name="Calibri", bold=True, color="4DB8A4", size=16)
    subtitle_font = Font(name="Calibri", color="666666", size=10)
    data_font = Font(name="Calibri", size=10)
    thin_border = Border(
        left=Side(style="thin", color="DDDDDD"),
        right=Side(style="thin", color="DDDDDD"),
        top=Side(style="thin", color="DDDDDD"),
        bottom=Side(style="thin", color="DDDDDD"),
    )
    left = Alignment(horizontal="left", vertical="center")
    center = Alignment(horizontal="center", vertical="center")

    ts = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    company = (project_meta or {}).get("company_name") or _db_get_setting("company_name", "TÁCTICA INGENIERÍA")
    project_line = _project_line(project_meta)
    summary = _summarize_items(items)

    logo_rows = 0
    if logo_path and logo_path.exists():
        try:
            img = XLImage(str(logo_path))
            img.width = 90
            img.height = 52
            ws_summary.add_image(img, "A1")
            logo_rows = 4
        except Exception:
            logo_rows = 1

    row_start = max(logo_rows, 1)
    ws_summary.merge_cells(start_row=row_start, start_column=2, end_row=row_start, end_column=6)
    ws_summary.cell(row=row_start, column=2, value=company).font = title_font
    ws_summary.merge_cells(start_row=row_start + 1, start_column=2, end_row=row_start + 1, end_column=6)
    ws_summary.cell(row=row_start + 1, column=2, value=titulo).font = subtitle_font
    ws_summary.merge_cells(start_row=row_start + 2, start_column=2, end_row=row_start + 2, end_column=6)
    ws_summary.cell(row=row_start + 2, column=2, value=f"Fecha: {ts}").font = subtitle_font

    meta_row = row_start + 4
    meta = [
        ("Proyecto/Cliente", project_line or "—"),
        ("Registros", summary["count"]),
        ("Campos", len(summary["headers"])),
        ("Versión", f"v{REPORT_VERSION}"),
    ]
    for i, (k, v) in enumerate(meta):
        ws_summary.cell(row=meta_row + i, column=2, value=k).font = header_font
        ws_summary.cell(row=meta_row + i, column=2).fill = teal_fill
        ws_summary.cell(row=meta_row + i, column=2).alignment = center
        ws_summary.cell(row=meta_row + i, column=3, value=v).font = data_font
        ws_summary.cell(row=meta_row + i, column=3).alignment = left

    stats_row = meta_row + len(meta) + 2
    ws_summary.cell(row=stats_row, column=2, value="Resumen ejecutivo").font = header_font
    ws_summary.cell(row=stats_row, column=2).fill = teal_fill
    ws_summary.merge_cells(start_row=stats_row, start_column=2, end_row=stats_row, end_column=6)
    ws_summary.cell(row=stats_row + 1, column=2, value=f"Se procesaron {summary['count']} registros.").font = data_font

    row_cursor = stats_row + 2
    for k, v in list(summary["numeric"].items())[:3]:
        ws_summary.cell(row=row_cursor, column=2, value=k).font = data_font
        ws_summary.cell(row=row_cursor, column=3, value=f"min {v['min']:.2f} | max {v['max']:.2f} | avg {v['avg']:.2f}").font = data_font
        row_cursor += 1

    ws_summary.column_dimensions["A"].width = 4
    ws_summary.column_dimensions["B"].width = 22
    ws_summary.column_dimensions["C"].width = 54

    # Detalle
    if items:
        headers = summary["headers"] or list(items[0].keys())
        for col_idx, header in enumerate(headers, 1):
            cell = ws_detail.cell(row=1, column=col_idx, value=str(header).upper())
            cell.font = header_font
            cell.fill = teal_fill
            cell.alignment = center
            cell.border = thin_border

        for row_idx, item in enumerate(items, 2):
            for col_idx, key in enumerate(headers, 1):
                val = item.get(key, "")
                cell = ws_detail.cell(row=row_idx, column=col_idx, value=val)
                cell.font = data_font
                cell.border = thin_border
                cell.alignment = left
                if (row_idx % 2) == 0:
                    cell.fill = alt_fill

        last_col = get_column_letter(len(headers))
        last_row = len(items) + 1
        ws_detail.auto_filter.ref = f"A1:{last_col}{last_row}"
        ws_detail.freeze_panes = "A2"

        for col_idx in range(1, len(headers) + 1):
            col_letter = get_column_letter(col_idx)
            max_len = max(len(str(headers[col_idx - 1])),
                          *[len(str(items[r].get(headers[col_idx - 1], ""))) for r in range(len(items))])
            ws_detail.column_dimensions[col_letter].width = min(max_len + 4, 38)

        if incluir_graficos and len(items) > 1:
            _agregar_graficos_excel(wb, ws_detail, items, headers, 1)

    wb.save(str(out_path))
    logger.info(f"Excel generado: {out_path}")

def generar_reporte_word_profesional(
    items: List[dict],
    out_path: Path,
    titulo: str = "Reporte",
    subtitulo: str = "",
    logo_path: Optional[Path] = None,
    project_meta: Optional[dict] = None
):
    """Genera un DOCX profesional (portada + resumen + tablas)."""
    if not DOCX_OK:
        logger.warning("python-docx no disponible, Word no generado")
        return

    if not logo_path:
        logo_path = _resolve_logo_path()

    doc = Document()

    company = (project_meta or {}).get("company_name") or _db_get_setting("company_name", "TÁCTICA INGENIERÍA")
    project_line = _project_line(project_meta)
    ts = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    summary = _summarize_items(items)

    # Portada
    if logo_path and logo_path.exists():
        try:
            doc.add_picture(str(logo_path), width=Pt(80))
        except Exception:
            pass

    p = doc.add_paragraph()
    run = p.add_run(company + "\n")
    run.bold = True
    run.font.size = Pt(18)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(titulo)
    run2.bold = True
    run2.font.size = Pt(16)

    if subtitulo:
        doc.add_paragraph(subtitulo)
    if project_line:
        doc.add_paragraph(f"Proyecto/Cliente: {project_line}")
    doc.add_paragraph(f"Fecha: {ts}")
    doc.add_paragraph(f"Versión del reporte: v{REPORT_VERSION}")

    doc.add_page_break()

    # Resumen ejecutivo
    doc.add_heading("Resumen ejecutivo", level=1)
    summary_lines = [f"Se procesaron {summary['count']} registros con {len(summary['headers'])} campos."]
    if summary["numeric"]:
        for k, v in list(summary["numeric"].items())[:3]:
            summary_lines.append(
                f"{k}: min {v['min']:.2f} | max {v['max']:.2f} | promedio {v['avg']:.2f}"
            )
    doc.add_paragraph("\n".join(summary_lines))

    # Detalle
    doc.add_heading("Detalle", level=1)
    if items:
        keys = summary["headers"] or list(items[0].keys())
        if not keys:
            keys = ["dato"]
        table = doc.add_table(rows=1, cols=len(keys))
        hdr = table.rows[0].cells
        for i, k in enumerate(keys):
            hdr[i].text = str(k)
        for it in items:
            row = table.add_row().cells
            for i, k in enumerate(keys):
                v = it.get(k, "")
                row[i].text = "" if v is None else str(v)
    else:
        doc.add_paragraph("No se encontraron datos para mostrar.")

    # Conclusiones y recomendaciones
    doc.add_heading("Conclusiones y recomendaciones", level=1)
    for c in [
        "Validar los registros clave y documentar supuestos de cálculo.",
        "Revisar los valores extremos detectados en las columnas numéricas.",
        "Consolidar esta versión como línea base del proyecto.",
    ]:
        doc.add_paragraph(c, style="List Bullet")

    # Footer con página y versión
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = "Página "
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    fp._p.append(fld)
    fp.add_run(f" | v{REPORT_VERSION}")

    doc.save(str(out_path))
    logger.info(f"Word generado: {out_path}")

def _agregar_graficos_excel(wb, ws_data, items, headers, data_start_row):
    numeric_cols = []
    label_col = None
    for h in headers:
        sample_vals = [items[i].get(h) for i in range(min(3, len(items)))]
        if any(isinstance(v, (int, float)) for v in sample_vals):
            numeric_cols.append(h)
        elif label_col is None:
            label_col = h

    if not numeric_cols or not label_col:
        return

    ws_chart = wb.create_sheet("Dashboard")

    for i, num_col in enumerate(numeric_cols[:3]):
        col_idx_label = headers.index(label_col) + 1
        col_idx_data = headers.index(num_col) + 1

        chart = BarChart()
        chart.type = "col"
        chart.title = num_col.replace("_", " ").title()
        chart.style = 10

        data_ref = Reference(ws_data,
                             min_col=col_idx_data,
                             min_row=data_start_row,
                             max_row=data_start_row + len(items))
        cat_ref = Reference(ws_data,
                            min_col=col_idx_label,
                            min_row=data_start_row + 1,
                            max_row=data_start_row + len(items))

        chart.add_data(data_ref, titles_from_data=True)
        chart.set_categories(cat_ref)
        chart.shape = 4
        chart.width = 18
        chart.height = 12

        series = chart.series[0]
        series.graphicalProperties.solidFill = "4DB8A4"

        ws_chart.add_chart(chart, f"A{1 + i * 16}")


def generar_grafico_bi(data: List[dict], campo_x: str, campo_y: str, titulo: str, out_path: Path, tipo: str = "bar"):
    if not MATPLOTLIB_OK:
        return None

    labels = [str(d.get(campo_x, "")) for d in data]
    values = []
    for d in data:
        v = d.get(campo_y, 0)
        try:
            values.append(float(v) if v else 0)
        except (ValueError, TypeError):
            values.append(0)

    fig, ax = plt.subplots(figsize=(10, 5))
    if tipo == "bar":
        ax.bar(labels, values)
        ax.set_ylabel(campo_y.replace("_", " ").title())
    ax.set_title(titulo)
    plt.tight_layout()
    plt.savefig(str(out_path), dpi=150, bbox_inches="tight")
    plt.close(fig)
    logger.info(f"GrÃ¡fico generado: {out_path}")
    return out_path


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  DATOS DE SIMULACIÃ“N
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

SIMUL_MEDICION = [
    {"elemento": "Viga V-1", "medida": 6.50, "unidad": "m", "observacion": "Eje A-B, piso 2"},
    {"elemento": "Columna C-3", "medida": 3.20, "unidad": "m", "observacion": "Altura libre"},
    {"elemento": "Losa L-1", "medida": 28.50, "unidad": "mÂ²", "observacion": "Entrepiso nivel +3.20"},
    {"elemento": "Muro M-2", "medida": 12.80, "unidad": "m", "observacion": "PerÃ­metro fachada norte"},
    {"elemento": "Aleta 3", "medida": 8.85, "unidad": "m", "observacion": "Puente vehicular"},
]

SIMUL_ESTRUCTURAL = [
    {"elemento": "Columna C-1", "tipo": "columna", "seccion": "40x40 cm", "longitud_m": 3.20,
     "area_m2": 0.16, "refuerzo": "8âˆ…20mm + Eâˆ…10mm c/15cm", "ubicacion": "Eje A-1", "observacion": "f'c=28MPa"},
    {"elemento": "Viga V-201", "tipo": "viga", "seccion": "30x50 cm", "longitud_m": 6.50,
     "area_m2": 0.15, "refuerzo": "4âˆ…20mm + 2âˆ…16mm + Eâˆ…10mm c/20cm", "ubicacion": "Eje 1 A-D", "observacion": "Piso 2"},
    {"elemento": "Zapata Z-1", "tipo": "zapata", "seccion": "180x180x40 cm", "longitud_m": 1.80,
     "area_m2": 3.24, "refuerzo": "âˆ…16mm c/15cm ambos sentidos", "ubicacion": "Eje A-1", "observacion": "Profundidad -1.50m"},
]

SIMUL_CARTILLA = [
    {"elemento": "Viga V-1", "marca": "1", "diametro_mm": 20, "cantidad": 4, "longitud_m": 6.80,
     "peso_unit_kg": 2.47, "peso_total_kg": 67.18, "forma": "Recta", "separacion_cm": "", "observacion": "Refuerzo inferior"},
]

SIMUL_PRESUPUESTO = [
    {"item": "1.1", "descripcion": "ExcavaciÃ³n mecÃ¡nica", "unidad": "mÂ³", "cantidad": 120,
     "valor_unitario": 35000, "valor_total": 4200000, "capitulo": "CimentaciÃ³n"},
]


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  ENDPOINTS â€” RENDER DE PLANOS (PDF -> PNG) PARA MEDICIÃ“N ROBUSTA
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.post("/api/v1/plan/render")
async def plan_render(
    file: UploadFile = File(...),
    page: int = Form(1),
    dpi: int = Form(150),
    project_id: Optional[str] = Form(None),
):
    contents = await file.read()
    ts = _report_timestamp()
    project_meta = _get_project_meta(project_id)
    filename = (file.filename or "plan").lower()

    # save original
    suffix = Path(file.filename or "").suffix
    if not suffix:
        suffix = ".pdf" if (file.content_type or "").lower() == "application/pdf" else ".png"
    orig_path = UPLOAD_DIR / f"asset_{ts}{suffix}"
    orig_path.write_bytes(contents)

    try:
        if filename.endswith(".pdf"):
            if fitz is None:
                raise HTTPException(500, "PyMuPDF no instalado. Instala: pip install PyMuPDF")
            doc = fitz.open(stream=contents, filetype="pdf")
            p = max(1, min(page, doc.page_count))
            zoom = max(72, min(dpi, 400)) / 72.0
            pix = doc.load_page(p-1).get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            out_name = f"render_{ts}_p{p}.png"
            out_path = OUTPUT_DIR / out_name
            out_path.write_bytes(pix.tobytes("png"))
            width_px, height_px = pix.width, pix.height
            page_count = doc.page_count
            doc.close()
            eid = _db_insert_event(project_id, "plan_rendered", {
                "filename": file.filename,
                "page": p,
                "dpi": dpi,
                "image": f"/outputs/{out_name}",
                "width_px": width_px,
                "height_px": height_px
            })
            return {
                "image_url": f"/outputs/{out_name}",
                "width_px": width_px,
                "height_px": height_px,
                "page": p,
                "page_count": page_count,
                "asset_saved": str(orig_path.name),
                "event_id": eid
            }

        # image passthrough
        img = PILImage.open(io.BytesIO(contents)).convert("RGB")
        out_name = f"render_{ts}.png"
        out_path = OUTPUT_DIR / out_name
        img.save(out_path, format="PNG")
        eid = _db_insert_event(project_id, "plan_rendered", {
            "filename": file.filename,
            "image": f"/outputs/{out_name}",
            "width_px": img.width,
            "height_px": img.height
        })
        return {
            "image_url": f"/outputs/{out_name}",
            "width_px": img.width,
            "height_px": img.height,
            "page": 1,
            "page_count": 1,
            "asset_saved": str(orig_path.name),
            "event_id": eid
        }
    except Exception as e:
        raise HTTPException(500, f"No se pudo renderizar: {str(e)}")

#  ENDPOINTS â€” ANÃLISIS GENERAL
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.post("/api/v1/analyze-and-report")
async def analyze_and_report(
    file: UploadFile = File(...),
    query: str = Form("Extrae las medidas"),
    mode: str = Form("measure"),
    project_id: Optional[str] = Form(None),
):
    contents = await file.read()
    ts = _report_timestamp()
    mode = normalize_mode(mode)
    project_meta = _get_project_meta(project_id)

    try:
        prompt_map = {
            "measure": PROMPT_MEDICION,
            "struct": PROMPT_ESTRUCTURAL,
            "rebar": PROMPT_CARTILLA_HIERRO,
            "budget": PROMPT_PRESUPUESTO,
        }
        prompt = prompt_map.get(mode, PROMPT_MEDICION)
        if query and query != "Extrae las medidas":
            prompt = f"{prompt}\n\nInstrucciÃ³n adicional del usuario: {query}"

        images = file_to_images(contents, file.filename)
        images_b64 = [image_to_base64(img) for img in images]

        if SIMULATION:
            simul_map = {
                "measure": SIMUL_MEDICION,
                "struct": SIMUL_ESTRUCTURAL,
                "rebar": SIMUL_CARTILLA,
                "budget": SIMUL_PRESUPUESTO,
            }
            items = simul_map.get(mode, SIMUL_MEDICION)
        else:
            raw = call_gpt4o_vision(prompt, images_b64)
            items = parse_json_from_response(raw)

        if not items:
            return JSONResponse({"items": [], "files": {}, "message": "No se encontraron datos en el anÃ¡lisis."})

        titulo_map = {
            "measure": "Reporte de Mediciones",
            "struct": "AnÃ¡lisis Estructural",
            "rebar": "Cartilla de Hierro â€” Despiece de Refuerzo",
            "budget": "AnÃ¡lisis de Presupuesto",
        }
        titulo = titulo_map.get(mode, "Reporte")

        files = {}

        pdf_name = f"{stem}.pdf"
        generar_reporte_pdf_profesional(
            items, OUTPUT_DIR / pdf_name,
            titulo=titulo,
            subtitulo=f"Archivo: {file.filename}",
            logo_path=_resolve_logo_path(),
            project_meta=project_meta,
        )
        files["pdf"] = f"/exports/{pdf_name}"

        xlsx_name = f"{stem}.xlsx"
        generar_reporte_excel_profesional(
            items, OUTPUT_DIR / xlsx_name,
            titulo=titulo,
            hoja="Detalle",
            logo_path=_resolve_logo_path(),
            incluir_graficos=(mode == "budget"),
            project_meta=project_meta,
        )
        files["excel"] = f"/exports/{xlsx_name}"

        if pd:
            csv_name = f"{stem}.csv"
            pd.DataFrame(items).to_csv(OUTPUT_DIR / csv_name, index=False)
            files["csv"] = f"/exports/{csv_name}"

        # Word (docx) opcional
        if DOCX_OK:
            docx_name = f"{stem}.docx"
            generar_reporte_word_profesional(
                items, OUTPUT_DIR / docx_name,
                titulo=titulo,
                subtitulo=f"Archivo: {file.filename}",
                logo_path=_resolve_logo_path(),
                project_meta=project_meta,
            )
            files["word"] = f"/exports/{docx_name}"

        # Guardar evento para trazabilidad / reportes bajo demanda
        analysis_id = _db_insert_event(project_id, "ai_analysis", {
            "mode": mode,
            "filename": file.filename,
            "query": query,
            "items": items,
            "files": files
        })

        return {"items": items, "files": files, "mode": mode, "count": len(items), "analysis_id": analysis_id}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en analyze: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  ENDPOINTS â€” CHAT IA
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.post("/api/v1/chat")
async def chat_ia(
    message: str = Form(...),
    mode: str = Form("measure"),
    file: Optional[UploadFile] = File(None),
    country: Optional[str] = Form(None),
    standard: Optional[str] = Form(None),
    project_id: Optional[str] = Form(None),
):
    """
    Endpoint de chat IA con manejo robusto de errores.
      500 -> falta la API key o error interno
      503 -> no se pudo conectar con OpenAI (red / firewall / proxy)
      504 -> timeout hacia OpenAI
      429 -> rate limit de OpenAI
      502 -> OpenAI respondio con error HTTP
      401 -> API key invalida
    """
    mode = normalize_mode(mode)

    # Validaciones tempranas
    if not OPENAI_API_KEY or not OPENAI_API_KEY.strip():
        logger.error("OPENAI_API_KEY no configurada")
        raise HTTPException(
            status_code=500,
            detail="OPENAI_API_KEY no configurada. Agregala en el archivo .env o en las variables de entorno del servidor."
        )

    if openai is None:
        raise HTTPException(
            status_code=503,
            detail="SDK de OpenAI no instalado. Ejecuta: pip install openai httpx"
        )

    # Modo demo / simulacion
    simulate = OPENAI_API_KEY.strip().upper() == "SIMULACION"
    if simulate:
        _db_insert_event(project_id, "chat", {"mode": mode, "message": message, "country": country, "standard": standard})
        return {
            "response": (
                f"[Modo demo] Recibi tu consulta: '{message}'. "
                "Para respuestas reales configura OPENAI_API_KEY con una clave valida."
            ),
            "mode": mode,
            "simulation": True,
        }

    # Contexto normativo
    ctx_norm = ""
    if country or standard:
        ctx_norm = f"\nContexto normativo: pais={country or 'N/A'} | norma={standard or 'N/A'}."

    system_map = {
        "measure": SYSTEM_PROMPT_BASE + "\nEstas en MEDICION." + ctx_norm,
        "struct":  SYSTEM_PROMPT_BASE + "\nEstas en ESTRUCTURAL." + ctx_norm,
        "rebar":   SYSTEM_PROMPT_BASE + "\nEstas en CARTILLA DE HIERRO." + ctx_norm,
        "budget":  SYSTEM_PROMPT_BASE + "\nEstas en PRESUPUESTOS." + ctx_norm,
    }
    system = system_map.get(mode, SYSTEM_PROMPT_BASE + ctx_norm)

    # Llamada a la IA
    # Los errores OAIConnectionError/OAITimeoutError/etc. se propagan desde
    # call_gpt4o_vision / call_gpt4o_text ya convertidos en HTTPException con
    # status 503/429/504/502, asi que solo necesitamos capturar lo inesperado aqui.
    try:
        if file and file.filename:
            contents = await file.read()
            lower_fn = (file.filename or "").lower()

            if lower_fn.endswith((".xlsx", ".xls")):
                # Excel -> texto estructurado (no imagen)
                data_rows = parse_excel_file(contents, file.filename)
                tabla_txt = json.dumps(data_rows[:80], ensure_ascii=False, default=str)
                augmented_msg = message + "\n\n[EXCEL: " + file.filename + "]\n" + tabla_txt
                response = call_gpt4o_text(augmented_msg, system=system)

            elif lower_fn.endswith(".csv"):
                if pd:
                    df_csv = pd.read_csv(io.BytesIO(contents))
                    tabla_txt = df_csv.head(80).to_csv(index=False)
                else:
                    tabla_txt = contents.decode("utf-8", errors="replace")[:4000]
                augmented_msg = message + "\n\n[CSV: " + file.filename + "]\n" + tabla_txt
                response = call_gpt4o_text(augmented_msg, system=system)

            else:
                # PDF e imagenes -> GPT-4o Vision
                images = file_to_images(contents, file.filename)
                images_b64 = [image_to_base64(img) for img in images]
                response = call_gpt4o_vision(message, images_b64, system=system)

        else:
            response = call_gpt4o_text(message, system=system)

        # Guardar evento en DB
        _db_insert_event(project_id, "chat", {
            "mode": mode,
            "message": message,
            "country": country,
            "standard": standard,
            "has_file": bool(file and file.filename),
        })

        return {
            "response": response or "Sin respuesta del modelo.",
            "mode": mode,
            "context": {"country": country, "standard": standard},
        }

    except HTTPException:
        # Re-propaga los 401/503/504/429/502 que vienen de call_gpt4o_*
        raise

    except Exception as e:
        # Ultimo recurso: errores no previstos (lectura de archivo, DB, etc.)
        logger.error(f"Chat error inesperado: {traceback.format_exc()}")
        raise HTTPException(
            status_code=500,
            detail=f"Error interno en /api/v1/chat: {str(e)}"
        )

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  ENDPOINTS â€” ESTRUCTURAL: VerificaciÃ³n cartilla vs plano
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.post("/api/v1/struct/verify")
async def struct_verify(
    plan: UploadFile = File(...),
    cartilla: UploadFile = File(...),
    country: str = Form("CO"),
    standard: str = Form("NSR-10"),
    project_id: Optional[str] = Form(None),
):
    country = (country or "CO").strip().upper()
    standard = (standard or "NSR-10").strip()

    plan_bytes = await plan.read()
    cart_bytes = await cartilla.read()
    ts = _report_timestamp()
    project_meta = _get_project_meta(project_id)

    # Convert inputs to images if needed
    def _to_images(b, name):
        lower = (name or "").lower()
        if lower.endswith((".xlsx", ".xls", ".csv")):
            # tabla: parse excel/csv a items determinÃ­sticos
            if lower.endswith(".csv"):
                if not pd:
                    raise HTTPException(500, "pandas requerido para CSV")
                df = pd.read_csv(io.BytesIO(b))
                return df.to_dict(orient="records")
            return parse_excel_file(b, name)
        imgs = file_to_images(b, name)
        return [image_to_base64(img) for img in imgs]

    # IA: anÃ¡lisis combinado (simple y efectivo)
    if SIMULATION:
        findings = [
            {"elemento": "Viga Eje B", "hallazgo": "DiÃ¡metro inconsistente entre plano (Ã˜16) y cartilla (Ã˜12)", "severidad": "alta"},
            {"elemento": "Columna C-3", "hallazgo": "SeparaciÃ³n de estribos no visible con claridad en plano", "severidad": "media"},
        ]
        summary = "Se detectaron 2 discrepancias relevantes. RecomendaciÃ³n: validar viga Eje B y aclarar estribos en C-3."
    else:
        plan_imgs_b64 = _to_images(plan_bytes, plan.filename)
        cart_data = _to_images(cart_bytes, cartilla.filename)

        prompt = f"""Tienes un plano estructural y una cartilla de hierro.
Objetivo: correlacionar y verificar coherencia entre ambos bajo contexto normativo paÃ­s={country}, norma={standard}.

Devuelve un JSON con:
{{
  "summary": "...",
  "findings": [
    {{"elemento":"...","hallazgo":"...","severidad":"alta|media|baja","evidencia":"..."}}
  ]
}}

SÃ© conservador: si no estÃ¡s seguro, marca severidad media/baja y explÃ­citalo.
"""

        # Si cartilla fue parseada (tabla), mandamos como texto; si no, mandamos imÃ¡genes.
        if isinstance(cart_data, list) and cart_data and isinstance(cart_data[0], dict):
            prompt2 = prompt + "\n\nCARTILLA (tabla JSON):\n" + json.dumps(cart_data[:200], ensure_ascii=False)
            raw = call_gpt4o_vision(prompt2, plan_imgs_b64, system=SYSTEM_PROMPT_BASE)
        else:
            # ambos como imÃ¡genes
            images_b64 = list(plan_imgs_b64)
            if isinstance(cart_data, list):
                images_b64.extend(cart_data)
            raw = call_gpt4o_vision(prompt, images_b64, system=SYSTEM_PROMPT_BASE)

        try:
            obj = json.loads(re.sub(r"```json|```", "", (raw or "").strip()))
        except Exception:
            obj = {"summary": raw or "Sin respuesta", "findings": []}

        summary = obj.get("summary", "")
        findings = obj.get("findings", [])

    # Guardar evento
    analysis_id = _db_insert_event(project_id, "struct_verified", {
        "country": country, "standard": standard,
        "plan_filename": plan.filename,
        "cartilla_filename": cartilla.filename,
        "summary": summary,
        "findings": findings
    })

    # Reportes
    files = {}
    stem = Path(_report_filename(ts, "pdf")).stem
    titulo = "Informe Estructural â€” VerificaciÃ³n Cartilla vs Plano"
    subt = f"{country} Â· {standard} | Plano: {plan.filename} | Cartilla: {cartilla.filename}"

    pdf_name = _report_filename(ts, "pdf")
    stem = Path(pdf_name).stem
    generar_reporte_pdf_profesional(findings, OUTPUT_DIR / pdf_name, titulo=titulo, subtitulo=subt, project_meta=project_meta)
    files["pdf"] = f"/exports/{pdf_name}"

    xlsx_name = f"{stem}.xlsx"
    generar_reporte_excel_profesional(findings, OUTPUT_DIR / xlsx_name, titulo=titulo, incluir_graficos=False, project_meta=project_meta)
    files["excel"] = f"/exports/{xlsx_name}"

    if DOCX_OK:
        docx_name = f"{stem}.docx"
        generar_reporte_word_profesional(findings, OUTPUT_DIR / docx_name, titulo=titulo, subtitulo=subt, project_meta=project_meta)
        files["word"] = f"/exports/{docx_name}"

    _db_insert_event(project_id, "report_generated", {"mode": "struct", "analysis_id": analysis_id, "files": files})

    return {"summary": summary, "findings": findings, "files": files, "analysis_id": analysis_id, "mode": "struct"}

#  ENDPOINTS â€” PRESUPUESTOS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.post("/api/v1/budget/upload")
async def budget_upload(
    file: UploadFile = File(...),
    tipo: str = Form("planned"),
    project_id: Optional[str] = Form(None),
):
    contents = await file.read()
    filename = file.filename.lower()

    try:
        if filename.endswith((".xlsx", ".xls")):
            data = parse_excel_file(contents, file.filename)
        elif filename.endswith(".csv"):
            if pd:
                df = pd.read_csv(io.BytesIO(contents))
                data = df.to_dict(orient="records")
            else:
                raise HTTPException(500, "pandas no instalado para CSV")
        else:
            raise HTTPException(400, f"Formato no soportado: {file.filename}. Usa .xlsx, .xls o .csv")

        save_path = UPLOAD_DIR / f"budget_{tipo}_{get_timestamp()}{Path(file.filename).suffix}"
        save_path.write_bytes(contents)

        # âœ… FIX: project_id puede ser None sin error de DB
        _db_insert_event(project_id, "budget_loaded", {
            "tipo": tipo, "filename": file.filename, "rows": len(data)
        })
        return {
            "status": "ok",
            "tipo": tipo,
            "filename": file.filename,
            "rows": len(data),
            "columns": list(data[0].keys()) if data else [],
            "data": data[:100],
            "saved_path": str(save_path.name),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Budget upload error: {traceback.format_exc()}")
        raise HTTPException(500, f"Error al procesar presupuesto: {str(e)}")


@app.post("/api/v1/budget/compare")
async def budget_compare(
    planned: str = Form(...),
    executed: str = Form(...),
    project_id: Optional[str] = Form(None),
):
    ts = _report_timestamp()

    try:
        planned_data = json.loads(planned)
        executed_data = json.loads(executed)

        comparison = []
        for p in planned_data:
            desc = p.get("descripcion", p.get("item", str(p)))
            planned_val = float(p.get("valor_total", p.get("total", 0)) or 0)

            exec_val = 0
            for e in executed_data:
                e_desc = e.get("descripcion", e.get("item", ""))
                if str(e_desc).strip().lower() == str(desc).strip().lower():
                    exec_val = float(e.get("valor_total", e.get("total", 0)) or 0)
                    break

            desviacion = exec_val - planned_val
            pct = (desviacion / planned_val * 100) if planned_val else 0

            comparison.append({
                "item": desc,
                "presupuestado": planned_val,
                "ejecutado": exec_val,
                "desviacion": desviacion,
                "desviacion_pct": round(pct, 1),
            })

        total_planned = sum(c["presupuestado"] for c in comparison)
        total_executed = sum(c["ejecutado"] for c in comparison)
        total_desv = total_executed - total_planned
        total_pct = (total_desv / total_planned * 100) if total_planned else 0

        chart_path = None
        if MATPLOTLIB_OK and comparison:
            chart_path = generar_grafico_bi(
                comparison, "item", "presupuestado",
                "Presupuestado vs Ejecutado",
                OUTPUT_DIR / f"chart_budget_{ts}.png",
                tipo="bar"
            )

        xlsx_name = _report_filename(ts, "xlsx")
        stem = Path(xlsx_name).stem
        generar_reporte_excel_profesional(
            comparison, OUTPUT_DIR / xlsx_name,
            titulo="Comparativo Presupuesto vs Ejecutado",
            hoja="Detalle",
            incluir_graficos=True,
            project_meta=project_meta,
        )

        pdf_name = f"{stem}.pdf"
        generar_reporte_pdf_profesional(
            comparison, OUTPUT_DIR / pdf_name,
            titulo="Informe Presupuesto vs Ejecutado",
            subtitulo=f"Total presupuestado: ${total_planned:,.0f} | Ejecutado: ${total_executed:,.0f}",
            project_meta=project_meta,
        )

        files = {"excel": f"/exports/{xlsx_name}", "pdf": f"/exports/{pdf_name}"}
        if chart_path:
            files["chart"] = f"/exports/{chart_path.name}"

        analysis_id = _db_insert_event(project_id, "budget_compared", {"comparison": comparison, "totals": {"presupuestado": total_planned, "ejecutado": total_executed}})
        return {
            "comparison": comparison,
            "totals": {
                "presupuestado": total_planned,
                "ejecutado": total_executed,
                "desviacion": total_desv,
                "desviacion_pct": round(total_pct, 1),
            },
            "files": files,
            "analysis_id": analysis_id,
        }

    except Exception as e:
        logger.error(f"Budget compare error: {traceback.format_exc()}")
        raise HTTPException(500, str(e))


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  ENDPOINTS â€” CALIBRACIÃ“N
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.post("/api/v1/calibrate")
async def calibrate(
    file: UploadFile = File(...),
    known_distance_m: float = Form(...),
    px_distance: float = Form(...),
    project_id: Optional[str] = Form(None),
):
    """âœ… FIX: project_id puede ser None. CalibraciÃ³n funciona sin proyecto activo."""
    if px_distance <= 0 or known_distance_m <= 0:
        raise HTTPException(400, "Distancias deben ser positivas")

    px_per_meter = px_distance / known_distance_m
    m_per_px = 1.0 / px_per_meter

    try:
        contents = await file.read()
        suffix = Path(file.filename or "plan").suffix or ".bin"
        save_path = UPLOAD_DIR / f"plan_calibrated_{get_timestamp()}{suffix}"
        save_path.write_bytes(contents)
        saved_name = str(save_path.name)
    except Exception as e:
        logger.warning(f"No pude guardar el archivo de calibraciÃ³n: {e}")
        saved_name = ""

    event_id = _db_insert_event(project_id, "calibration", {
        "filename": getattr(file, "filename", ""),
        "known_distance_m": known_distance_m,
        "px_distance": px_distance,
        "px_per_meter": round(px_per_meter, 6),
        "m_per_px": round(m_per_px, 8),
    })

    return {
        "px_per_meter": round(px_per_meter, 4),
        "m_per_px": round(m_per_px, 6),
        "known_distance_m": known_distance_m,
        "px_distance": px_distance,
        "plan_saved": saved_name,
        "event_id": event_id,
    }


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  ENDPOINTS â€” REPORTES BAJO DEMANDA
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.post("/api/v1/report/generate")
async def generate_report(
    request: Request,
    data: Optional[str] = Form(None),
    format: str = Form("pdf"),
    titulo: str = Form("Reporte TÃ¡ctica"),
    subtitulo: str = Form(""),
    mode: str = Form("measure"),
    project_id: Optional[str] = Form(None),
    analysis_id: Optional[str] = Form(None),
):
    ts = _report_timestamp()
    mode = normalize_mode(mode)

    # 1) Obtener items (FormData o JSON o analysis_id)
    items = None
    ct = (request.headers.get("content-type") or "").lower()

    if data:
        try:
            items = json.loads(data)
        except json.JSONDecodeError:
            raise HTTPException(400, "JSON invÃ¡lido en 'data'")
    elif "application/json" in ct:
        body = await request.json()
        format = body.get("format", format)
        titulo = body.get("titulo", titulo)
        subtitulo = body.get("subtitulo", subtitulo)
        mode = normalize_mode(body.get("mode", mode))
        project_id = body.get("project_id", project_id)
        analysis_id = body.get("analysis_id", analysis_id)
        items = body.get("items") or body.get("data") or None

        # si items viene como string JSON
        if isinstance(items, str):
            try:
                items = json.loads(items)
            except json.JSONDecodeError:
                items = None

    # 2) Si no hay items, intentar desde analysis/event
    if items is None and analysis_id:
        ev = _db_get_event(analysis_id)
        if ev and isinstance(ev.get("payload", {}).get("items"), list):
            items = ev["payload"]["items"]

    if items is None:
        raise HTTPException(400, "No hay 'data/items' ni 'analysis_id' para generar el reporte.")

    if not isinstance(items, list):
        items = [items]

    fmt = (format or "pdf").lower()
    # alias
    if fmt == "xlsx":
        fmt = "excel"
    if fmt == "docx":
        fmt = "word"
    if fmt == "all":
        fmt = "all"

    # Project meta
    project_meta = _get_project_meta(project_id)

    files = {}
    stem = Path(_report_filename(ts, "pdf")).stem

    def _mk_titles():
        # si no hay subtitulo, mete algo Ãºtil
        st = subtitulo or (f"Proyecto: {project_meta.get('project_name','')}" if project_meta.get("project_name") else "")
        return titulo, st

    t, st = _mk_titles()

    if fmt in ("pdf", "both", "all"):
        pdf_name = f"{stem}.pdf"
        generar_reporte_pdf_profesional(items, OUTPUT_DIR / pdf_name, titulo=t, subtitulo=st, project_meta=project_meta)
        files["pdf"] = f"/exports/{pdf_name}"

    if fmt in ("excel", "both", "all"):
        xlsx_name = f"{stem}.xlsx"
        generar_reporte_excel_profesional(
            items, OUTPUT_DIR / xlsx_name,
            titulo=t,
            incluir_graficos=(mode == "budget"),
            project_meta=project_meta,
        )
        files["excel"] = f"/exports/{xlsx_name}"

    if fmt in ("word", "all"):
        if not DOCX_OK:
            raise HTTPException(500, "python-docx no instalado. Instala: pip install python-docx")
        docx_name = f"{stem}.docx"
        generar_reporte_word_profesional(items, OUTPUT_DIR / docx_name, titulo=t, subtitulo=st, project_meta=project_meta)
        files["word"] = f"/exports/{docx_name}"

    report_id = _new_id("rep")
    conn = _db_connect()
    conn.execute(
        "INSERT INTO reports(id, project_id, module, title, subtitle, files_json, source_event_id, created_at) VALUES(?,?,?,?,?,?,?,?)",
        (report_id, project_id, mode, t, st, json.dumps(files, ensure_ascii=False), analysis_id, _now_iso())
    )
    conn.commit()
    conn.close()

    _db_insert_event(project_id, "report_generated", {"mode": mode, "format": format, "files": files, "report_id": report_id})

    return {"files": files, "count": len(items), "report_id": report_id}



# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  ENDPOINTS â€” PROYECTOS (NÃºcleo del Proyecto)
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.post("/api/v1/projects")
async def create_project(
    name: str = Form(...),
    location: str = Form(""),
    currency: str = Form("COP"),
    budget_total: Optional[float] = Form(None),
):
    pid = _new_id("prj")
    now = _now_iso()
    conn = _db_connect()
    conn.execute(
        "INSERT INTO projects(id,name,location,currency,budget_total,created_at,updated_at) VALUES(?,?,?,?,?,?,?)",
        (pid, name.strip(), location.strip(), (currency or "COP").strip().upper(), budget_total, now, now)
    )
    conn.commit()
    conn.close()
    _db_insert_event(pid, "project_created", {"name": name, "location": location, "currency": currency, "budget_total": budget_total})
    return {"id": pid, "created_at": now}

def _project_kpis(project_id: str) -> dict:
    conn = _db_connect()
    rows = conn.execute("SELECT type, COUNT(*) as c FROM events WHERE project_id=? GROUP BY type", (project_id,)).fetchall()
    conn.close()
    counts = {r["type"]: r["c"] for r in rows}

    def ratio(done, total):
        if total <= 0:
            return 0.0
        return max(0.0, min(1.0, done/total))

    plans = counts.get("plan_rendered", 0) + counts.get("plan_loaded", 0)
    measures = counts.get("measurement_created", 0)
    measure_progress = ratio(measures, max(1, plans))

    cartillas = counts.get("rebar_loaded", 0)
    struct_ok = counts.get("struct_verified", 0)
    struct_progress = ratio(struct_ok, max(1, cartillas))

    budget_loaded = counts.get("budget_loaded", 0)
    budget_done = counts.get("budget_compared", 0)
    budget_progress = ratio(budget_done, max(1, budget_loaded))

    return {
        "measure_progress": round(measure_progress, 3),
        "struct_progress": round(struct_progress, 3),
        "budget_progress": round(budget_progress, 3),
        "counts": counts
    }

@app.get("/api/v1/projects")
async def list_projects():
    conn = _db_connect()
    rows = conn.execute("SELECT * FROM projects ORDER BY updated_at DESC").fetchall()
    conn.close()
    projects = []
    for r in rows:
        projects.append({
            "id": r["id"],
            "name": r["name"],
            "location": r["location"],
            "currency": r["currency"],
            "budget_total": r["budget_total"],
            "created_at": r["created_at"],
            "updated_at": r["updated_at"],
            "kpis": _project_kpis(r["id"])
        })
    return {"projects": projects}

@app.get("/api/v1/projects/{project_id}")
async def get_project(project_id: str):
    conn = _db_connect()
    r = conn.execute("SELECT * FROM projects WHERE id=?", (project_id,)).fetchone()
    conn.close()
    if not r:
        raise HTTPException(404, "Proyecto no encontrado")
    return {
        "id": r["id"],
        "name": r["name"],
        "location": r["location"],
        "currency": r["currency"],
        "budget_total": r["budget_total"],
        "created_at": r["created_at"],
        "updated_at": r["updated_at"],
        "kpis": _project_kpis(r["id"])
    }

@app.post("/api/v1/projects/{project_id}/events")
async def add_event(project_id: str, request: Request):
    body = await request.json()
    type_ = body.get("type") or "event"
    payload = body.get("payload") or {}
    eid = _db_insert_event(project_id, type_, payload)
    # touch updated_at
    conn = _db_connect()
    conn.execute("UPDATE projects SET updated_at=? WHERE id=?", (_now_iso(), project_id))
    conn.commit()
    conn.close()
    return {"event_id": eid, "stored": True}

@app.get("/api/v1/projects/{project_id}/events")
async def list_events(project_id: str, limit: int = 100):
    conn = _db_connect()
    rows = conn.execute(
        "SELECT * FROM events WHERE project_id=? ORDER BY created_at DESC LIMIT ?",
        (project_id, min(max(limit,1), 500))
    ).fetchall()
    conn.close()
    return {"events": [
        {"id": r["id"], "type": r["type"], "payload": json.loads(r["payload_json"]), "created_at": r["created_at"]}
        for r in rows
    ]}

@app.get("/api/v1/projects/{project_id}/reports")
async def list_reports(project_id: str, limit: int = 50):
    conn = _db_connect()
    rows = conn.execute(
        "SELECT * FROM reports WHERE project_id=? ORDER BY created_at DESC LIMIT ?",
        (project_id, min(max(limit,1), 200))
    ).fetchall()
    conn.close()
    return {"reports": [
        {
            "id": r["id"],
            "module": r["module"],
            "title": r["title"],
            "subtitle": r["subtitle"],
            "files": json.loads(r["files_json"]),
            "source_event_id": r["source_event_id"],
            "created_at": r["created_at"]
        } for r in rows
    ]}


# ─────────────────────────────────────────────────────────────────────────────
#  ENDPOINT DE DIAGNÓSTICO DE RED (netcheck)
#  Útil para verificar si el servidor puede alcanzar api.openai.com
#  sin depender del frontend. Llama a GET /api/v1/netcheck
# ─────────────────────────────────────────────────────────────────────────────

@app.get("/api/v1/netcheck")
async def netcheck():
    """
    Verifica si el servidor puede abrir una conexion TCP al puerto 443 de api.openai.com.
    - ok: true  -> la conexion funciona, el problema es otro (API key, rate limit, etc.)
    - ok: false -> el trafico esta bloqueado (firewall, proxy, DNS, etc.)
    """
    import socket as _socket
    host = "api.openai.com"
    port = 443
    s = _socket.socket(_socket.AF_INET, _socket.SOCK_STREAM)
    s.settimeout(5)
    try:
        s.connect((host, port))
        s.close()
        return {
            "ok": True,
            "host": host,
            "port": port,
            "detail": f"Conexion TCP al puerto {port} de {host} exitosa. La red funciona.",
        }
    except Exception as e:
        return {
            "ok": False,
            "host": host,
            "port": port,
            "detail": f"Puerto {port} BLOQUEADO o filtrado hacia {host}: {str(e)}. "
                      "Revisa firewall, proxy o DNS. Configura HTTPS_PROXY si usas red corporativa.",
        }
    finally:
        try:
            s.close()
        except Exception:
            pass

#  HEALTH
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.get("/api/v1/health")
async def health():
    return {
        "status": "ok",
        "version": "3.1.0",
        "simulation": SIMULATION,
        "modules": {
            "openai": openai is not None,
            "fitz": fitz is not None,
            "pandas": pd is not None,
            "reportlab": REPORTLAB_OK,
            "openpyxl": OPENPYXL_OK,
            "matplotlib": MATPLOTLIB_OK,
        },
        "static_dir": str(STATIC_DIR),
        "static_exists": STATIC_DIR.exists(),
        "timestamp": datetime.datetime.now().isoformat(),
    }


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  RUTAS WEB (PWA + shortcuts)
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.get("/", include_in_schema=False)
async def root():
    # âœ… redirige a /static (verificando existencia, incluso si quedÃ³ *.html.html)
    if _resolve_static("index.html") is not None:
        return RedirectResponse(url="/static/index.html")
    if _resolve_static("viewer.html") is not None:
        return RedirectResponse(url="/static/viewer.html")
    raise HTTPException(404, "No se encontrÃ³ index.html ni viewer.html en /static")


@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    rel = _resolve_static("favicon.ico")
    if rel is not None:
        return FileResponse(str((STATIC_DIR / rel).resolve()))
    raise HTTPException(404, "favicon.ico no encontrado")


@app.get("/manifest.json", include_in_schema=False)
async def manifest():
    return JSONResponse({
        "name": "TÃ¡ctica Pro â€” Visor Inteligente",
        "short_name": "TÃ¡ctica Pro",
        "description": "MediciÃ³n, anÃ¡lisis estructural y presupuestos para ingenierÃ­a civil",
        "start_url": "/static/index.html",
        "display": "standalone",
        "orientation": "landscape",
        "background_color": "#060b10",
        "theme_color": "#4db8a4",
        "icons": [
            {"src": "/static/assets/icon-192.png", "sizes": "192x192", "type": "image/png"},
            {"src": "/static/assets/icon-512.png", "sizes": "512x512", "type": "image/png"},
        ]
    })


@app.get("/sw.js", include_in_schema=False)
async def service_worker():
    # âœ… antes lo estabas mandando como JSON; asÃ­ NO funciona un SW
    sw_content = """
self.addEventListener('install', e => { self.skipWaiting(); });
self.addEventListener('activate', e => { e.waitUntil(clients.claim()); });
self.addEventListener('fetch', e => {
  e.respondWith(fetch(e.request).catch(() => new Response('Offline', {status: 503})));
});
"""
    return Response(content=sw_content, media_type="application/javascript")


# âœ… Shortcuts limpios (sin duplicados)
@app.get("/viewer.html", include_in_schema=False)
async def viewer_page():
    return RedirectResponse("/static/viewer.html")

@app.get("/projects.html", include_in_schema=False)
async def projects_page():
    return RedirectResponse("/static/projects.html")

@app.get("/estructural.html", include_in_schema=False)
async def estructural_page():
    # en tu carpeta el archivo real es structural.html
    return RedirectResponse("/static/structural.html")

@app.get("/budget.html", include_in_schema=False)
async def budget_page():
    # en tu carpeta el archivo real es budget_advanced.html
    return RedirectResponse("/static/budget_advanced.html")

@app.get("/dashboard.html", include_in_schema=False)
async def dashboard_page():
    return RedirectResponse("/static/dashboard_panel.html")


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  INICIO DEL SERVIDOR
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

if __name__ == "__main__":
    print("â•" * 60)
    print("  TÃCTICA INGENIERÃA â€” Servidor v3.1")
    print(f"  Modo: {'SIMULACIÃ“N' if SIMULATION else 'PRODUCCIÃ“N (GPT-4o)'}")
    print(f"  Static: {STATIC_DIR} | existe={STATIC_DIR.exists()}")
    print("â•" * 60)

    # ðŸ”¥ Tip pro: si en Windows se te mueve el working dir, esto lo estabiliza:
    # uvicorn tactica_profesional:app --reload --port 8000 --app-dir "C:\\tactica-profesional-completa"
    uvicorn.run(
        "tactica_profesional:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info",
    )
